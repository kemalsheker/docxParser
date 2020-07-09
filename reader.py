from docx import Document
import pandas as pd
document = Document("readfrom.docx")
tables = list(document.tables)

k = 4
it = 0

for i in tables:
    print(i.cell(0, 0).text)
    components = {'Component': [ i.cell(0,0).text ]}
    df = pd.DataFrame(components, columns=['Component'])
    pd.DataFrame(df).to_csv('inputdeneme.csv', mode='a')
    print(i.cell(5, 2).tables)
    iteration = len(i.rows) - 3
    k = 4
    while k < iteration:
        df = [['' for a in range(len(i.cell(k, 2).tables[0].columns))] for b in range(len(i.cell(k, 2).tables[0].rows))]
        for a, row in enumerate(i.cell(k, 2).tables[0].rows):
            for b, cell in enumerate(row.cells):
                if cell.text == "Input name" or cell.text == "Source":
                    df[a][b] = cell.text
                if b == 0 or b == 2:
                    df[a][b] = cell.text
        print(df)
        k = k + 1
        pd.DataFrame(df).to_csv('inputdeneme.csv', mode='a')



   # pd.DataFrame(df).to_csv('outputdeneme.csv', mode='a')




""""
print(len(tables))

for index, table in enumerate(document.tables):
    df = [['' for i in range(len(table.columns))] for j in range(len(table.rows))]
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if cell.text == "Input name" or cell.text == "Source":
                df[i][j] = cell.text
            if j == 0 or j == 2:
                df[i][j] = cell.text

    pd.DataFrame(df).to_csv('outputdeneme.csv', mode='a')
"""
